'輔助PANDAS工具'
from pathlib import Path
import logging
logger = logging.getLogger(Path(__file__).stem)

class 使用者要求更新且線上資料較線下多(Exception): pass
class 使用者要求覆寫(Exception): pass

def 增加自動更新資料功能(更新頻率='每月十日之前'):
    '如將參數「更新」設為True，則強制更新該批資料'
    import logging
    from pathlib import Path
    from diskcache import Cache
    from functools import wraps
    from zhongwen.date import 今日 
    cache = Cache(Path.home() / 'cache' / Path(__file__).stem)
    def _增加按期更新查詢結果功能(查詢資料):
        @wraps(查詢資料)
        def 查詢按期更新資料(*args, 更新=False,**kargs):
            if 更新頻率=='每月十日之前':
                if 今日().day <= 10 or 更新:
                    df = 查詢資料(*args, **kargs)
                    cache.set(查詢資料.__name__, df)
                    return df
            return cache.read(查詢資料.__name__)
        return 查詢按期更新資料
    return _增加按期更新查詢結果功能

class 應更新(Exception): pass

def 增加按期更新查詢結果功能(更新頻率='每月十日之前'):
    '如將參數「更新」設為True，則強制更新該批資料'
    import logging
    from pathlib import Path
    from diskcache import Cache
    from zhongwen.date import 今日 
    warn(f'【增加按期更新查詢結果功能】將廢棄，請使用【batch_data.增加定期更新】', DeprecationWarning, stacklevel=2)
    cache = Cache(Path.home() / 'cache' / Path(__file__).stem)
    def _增加按期更新查詢結果功能(查詢資料):
        from functools import wraps
        @wraps(查詢資料)
        def 查詢按期更新資料(*args, 更新=False,**kargs):
            try:
                if 更新 or (更新頻率=='每月十日之前' and 今日().day <= 10): raise 應更新()
                return cache.read(查詢資料.__name__)
            except (KeyError, 應更新) as e:
                df = 查詢資料(*args, **kargs)
                cache.set(查詢資料.__name__, df)
                return df
        return 查詢按期更新資料
    return _增加按期更新查詢結果功能

def 增加批次儲存(資料庫檔, 資料名稱, 批號欄名):
    from functools import wraps
    from sqlite3 import connect
    import pandas as pd
    import logging
    from warnings import warn
    warn(f'【增加批次緩存】將廢棄，請使用【batch_data.結果批次寫入】', DeprecationWarning, stacklevel=2)
    def 增加批次儲存功能(資料存取函數):
        @wraps(資料存取函數)
        def 批次資料存取(批號, *args, **kargs):
            df = 資料存取函數(批號, *args, **kargs)
            with connect(資料庫檔) as db: 
                批次寫入(df, 批號, 批號欄名, 資料名稱, db)
            return df
        return 批次資料存取
    return 增加批次儲存功能

def 增加批次緩存功能(資料庫檔, 資料名稱, 批號欄名):
    '如將參數「更新」設為True，則強制更新該批資料'
    from functools import wraps
    from sqlite3 import connect
    import pandas as pd
    import logging
    from warnings import warn
    warn(f'【增加批次緩存功能】將廢棄，請使用【batch_data.結果批次寫入】', DeprecationWarning, stacklevel=2)
    def _增加批次緩存功能(資料存取函數):
        @wraps(資料存取函數)
        def 批次緩存資料存取(批號, *args, 更新=False, 覆寫=False,**kargs):
            with connect(資料庫檔) as db: 
                try:
                    df0 = 批次讀取(批號, 批號欄名, 資料名稱, db) 
                    if 更新:
                        df1 = 資料存取函數(批號, *args, **kargs)
                        if 覆寫: raise 使用者要求覆寫()
                        if df1.shape[0] > df0.shape[0]:
                            logging.info(
                                f'目前抓取{批號}之{資料名稱}計{df1.shape[0]}筆，'
                                f'已較緩存資料庫{df0.shape[0]}筆多，爰更新資料庫！')    
                            raise 使用者要求更新且線上資料較線下多()
                    return df0 
                except (pd.errors.DatabaseError, 批號查無資料錯誤) as e:
                    logging.info(e)
                    df1 = 資料存取函數(批號, *args, **kargs)
                    批次寫入(df1, 批號, 批號欄名, 資料名稱, db, 更新)
                    return df1
                except (使用者要求覆寫, 使用者要求更新且線上資料較線下多) as e:
                    logging.info(e)
                    批次寫入(df1, 批號, 批號欄名, 資料名稱, db, 覆寫=True)
                    return df1
        return 批次緩存資料存取
    return _增加批次緩存功能

class 批號存在錯誤(Exception):
    def __init__(self, 批號, 批號欄名, 表格):
        self.批號 = 批號
        self.批號欄名 = 批號欄名
        self.表格 = 表格

    def __str__(self):
        return f'批號【{self.批號}】已存在，請指定覆寫=True！'

class 批號查無資料錯誤(Exception): pass

def 批次讀取(批號, 批號欄名, 表格, 資料庫, parse_dates=None):
    import pandas as pd
    from zhongwen.date import 是日期嗎
    # sql 的相等係單等號表示，而python為雙等號，易誤植
    sql = f'select * from {表格} where {批號欄名}={批號}' 
    if 是日期嗎(批號) or isinstance(批號, str):
        sql = f'select * from {表格} where {批號欄名}="{批號}"' 
    df = pd.read_sql_query(sql, 資料庫, parse_dates=parse_dates)
    if df.empty: 
        raise 批號查無資料錯誤(f'查無批號為【{批號}】資料！')
    return df

def 批次刪除(批號, 批號欄名, 表格, 資料庫):
    c = 資料庫.cursor()
    sql = f'delete from {表格} where {批號欄名}=="{批號}"'
    c.execute(sql)
    資料庫.commit()
    import logging
    logging.debug(f'批次刪除{批號}、{表格}……')

def 批次寫入(資料, 批號, 批號欄名, 表格, 資料庫, 覆寫=False):
    from warnings import warn
    import logging
    import pandas as pd
    import sqlite3
    import re
    warn(f'參數【覆寫】將廢棄。', DeprecationWarning, stacklevel=2)
    try:
        批次刪除(批號, 批號欄名, 表格, 資料庫)
    except sqlite3.OperationalError as e:
        logging.debug(f'批次刪除發生錯誤{e}')
    logging.debug(f'批次寫入{批號}、{表格}……')
    try:
        資料.to_sql(表格, 資料庫, if_exists='append')
    except sqlite3.OperationalError as e:
        df = pd.read_sql_query(f'select * from {表格} limit 1', 資料庫)
        寫入資料欄位 = 資料.columns.to_list()
        資料庫欄位 = df.columns.to_list()
        差異欄位 = set(寫入資料欄位) - set(資料庫欄位)
        print(差異欄位)
        cursor = 資料庫.cursor()
        for c in 差異欄位:
            alter_query = f'ALTER TABLE {表格} ADD COLUMN "{c}" INT'
            cursor.execute(alter_query)
        資料庫.commit()
        print('已新增差異欄位')
        資料.to_sql(表格, 資料庫, if_exists='append')
        # raise e
    logging.debug(f'寫入成功！')

def 載入時間序列(資料庫檔, 表格, 時間欄名, 最早時間=None, 實體欄名=None):
    '由實體及時間組成索引鍵'
    import pandas as pd
    import sqlite3
    with sqlite3.connect(資料庫檔) as c:
        sql = f"select * from {表格}"
        if 最早時間:
            sql += f" where {時間欄名} >= '{最早時間:%Y-%m-%d}' order by {時間欄名}"
        df = pd.read_sql_query(sql, c, index_col='index', parse_dates=時間欄名) 
        df[實體欄名] = df[實體欄名].astype(str)
        return df.set_index([實體欄名, 時間欄名])

def 可顯示(查詢資料函數):
    '裝飾查詢資料函數，指名參數設為【顯示=True】，即將查詢結果以 html 顯示。'
    from functools import wraps
    import matplotlib.pyplot as plt
    @wraps(查詢資料函數)
    def wrapper(*args, 顯示=False, 圖示=False, **kargs):
        try:
            display_rows = kargs['顯示筆數']
            real_columns = kargs['實數欄位']
            del kargs['顯示筆數']
            del kargs['實數欄位']
        except KeyError:
            display_rows = 100
            real_columns = []
        df = 查詢資料函數(*args, **kargs)
        if 顯示: show_html(df, 顯示筆數=display_rows, 實數欄位=real_columns)
        if 圖示: 
            plt.rcParams['font.sans-serif'] = ['Microsoft JhengHei']
            df.plot()
            plt.show()
        return df
    return wrapper

def 強調關鍵字(df, 欄位=[], 關鍵字=[]):
    for c in 欄位:
        for k in 關鍵字:
            df[c] = df[c].str.replace(k, f'<b style="background:pink">{k}</b>')
    return df

def 資料型態標準化(df):
    '依資料框欄位名稱自動將資料框之資料型態標準化'
    from zhongwen.date import 取日期
    from zhongwen.number import 轉數值
    for c in df.columns:
        p = f'.*日期'
        import re
        if re.match(p, c):
            df[c] = df[c].map(取日期)
        p = f'.*金額'
        if re.match(p, c):
            df[c] = df[c].map(轉數值)
    return df            

def 標準格式(整數欄位=[], 實數欄位=[], 百分比欄位=[], 日期欄位=[], 
             最大值顯著欄位=[], 隱藏欄位=[], 
             百分比漸層欄位=[], 
             採用民國日期格式=False
            ):
    import logging
    logging.debug(f'整數欄位：{整數欄位!r}')
    logging.debug(f'實數欄位：{實數欄位!r}')
    logging.debug(f'百分比欄位：{百分比欄位!r}')
    logging.debug(f'百分比漸層欄位：{百分比漸層欄位!r}')
    logging.debug(f'最大值顯著欄位：{最大值顯著欄位!r}')
    logging.debug(f'日期欄位：{日期欄位!r}')
    logging.debug(f'隱藏欄位：{隱藏欄位!r}')

    def formatter(style):
        style.map(lambda r:'text-align:right')
        if 整數欄位:
            style.format('{:,.0f}', subset=整數欄位)
        if 百分比欄位:
            style.format('{:,.2%}', subset=百分比欄位)
        if 實數欄位:
            style.format('{:,.2f}', subset=實數欄位)
        if 日期欄位:
            style.format('{:%Y%m%d}', subset=日期欄位)
            if 採用民國日期格式:
                from zhongwen.date import 民國日期
                style.format(民國日期, subset=日期欄位)
        if 最大值顯著欄位:
            # style.highlight_max(最大值顯著欄位)
            style.background_gradient(axis=0, cmap='RdYlGn'
                                     ,subset=最大值顯著欄位)
        if 百分比漸層欄位:
            style.background_gradient(axis=0, cmap='RdYlGn', vmax=1, vmin=-1
                                     ,subset=百分比漸層欄位
                                     )
        if 隱藏欄位:
            style.hide(隱藏欄位, axis=1) # hide index
        tr_hover = {
            'selector': 'tr:hover',
            'props':[('background-color', '#ffffb3'), ('border-style', 'dotted')]
        }
        style.set_table_styles([tr_hover], overwrite=False)


        return style
    return formatter

def show_html(df, 無格式=False
             ,整數欄位=[], 實數欄位=[]
             ,百分比欄位=[] ,日期欄位=[] ,隱藏欄位=[]
             ,百分比漸層欄位=[]
             ,最大值顯著欄位=[]
             ,顯示筆數=100, 採用民國日期格式=False, 標題=None
             ,傳回超文件內容=False
             ):
    from warnings import warn
    from pathlib import Path
    import pandas as pd
    import numpy as np
    import os

    if isinstance(df, set):
        df = pd.Series(list(df)).to_frame()

    if isinstance(df, list):
        df = pd.Series(df).to_frame()

    if isinstance(df, np.ndarray):
        df = pd.Series(df).to_frame()

    if isinstance(df, pd.Series):
        df = df.to_frame()

    if isinstance(df, pd.DataFrame):
        df = df.dropna(axis='columns', how='all')
        if not df.index.is_unique:
            df = df.reset_index(drop=True)

    if isinstance(df, pd.DataFrame) and not 無格式:
        try:
            df = df.head(顯示筆數)
            return 自動格式(df
                           ,整數欄位
                           ,實數欄位
                           ,百分比欄位
                           ,日期欄位
                           ,隱藏欄位
                           ,百分比漸層欄位
                           ,最大值顯著欄位
                           ,顯示筆數
                           ,採用民國日期格式
                           ,顯示=True
                           )
        except (ValueError, KeyError, TypeError) as e: 
            import traceback
            trace = traceback.format_exc()
            warn(f'發生例外：{trace}')
    html = Path.home() / 'TEMP' / 'output.html'
    df.to_html(html)
    os.system(f'start {html}')

def 自動格式(df, 整數欄位=[] ,實數欄位=[], 百分比欄位=[]
            ,日期欄位=[] ,隱藏欄位=[]
            ,百分比漸層欄位=[], 最大值顯著欄位=[], 不排序欄位=[]
            ,顯示筆數=100, 採用民國日期格式=False, 顯示=None
            ,除錯提示=False, 顯示提示=True
            ):
    from zhongwen.number import 轉數值
    import pandas as pd
    import numpy as np
    import re
    # df = df.fillna('')
    if 顯示筆數:
        df = df[:顯示筆數]
    df.columns = df.columns.fillna('unnamed')
    columns = df.columns
    for c in df.columns:
        try:
            pat = '^.*述|借貸$'
            if re.match(pat, c):
                continue
        except TypeError:
            logger.error(f'欄名{c}非字串格式')
            continue

        pat = '^.*(金額|專戶|淨股利|次數|損益|[毛淨營]利|累計|差異|評分|期末|負債|營收|年數|\(元\))|成本|支出|存入|現值|借券|餘額|借|貸$'
        if re.match(pat, c):
            try:
                if (np.issubclass_(df[c].dtype.type, np.integer)  or
                    df[c].dtype == float or df[c].dtype == 'float64'
                ) and not c in 隱藏欄位:
                    try:
                        整數欄位.append(c)
                    except AttributeError:
                        整數欄位 = [c]
            except:
                breakpoint()
                pass
        pat = '^現金轉換天數|估計每股配發現金|估計股利|(元/股)|股價|配息|配股|r方|每股盈餘|.*符合度|.*比|.*指數|每股.*$'
        if re.match(pat, c):
            if df[c].dtype in [float] and not c in 隱藏欄位 and not c in 百分比欄位 and not c in 整數欄位: 
                try:
                    實數欄位.append(c)
                except AttributeError:
                    實數欄位 = [c]
        pat = '^.*(漲跌幅|率|比例|趨勢|\(%\))$'
        if re.match(pat, c):
            if df[c].dtype == float and not c in 隱藏欄位: 
                try:
                    百分比欄位.append(c)
                except AttributeError:
                    百分比欄位 = [c]
        pat = '.*(日期|時期).*'
        if re.match(pat, c):
            if df[c].dtype == "datetime64[ns]" and not c in 隱藏欄位:
                try:
                    日期欄位.append(c)
                except AttributeError:
                    日期欄位 = [c]

    if 整數欄位: 整數欄位 = [*set(整數欄位)]
    if 實數欄位: 實數欄位 = [*set(實數欄位)]
    if 百分比欄位: 百分比欄位 = [*set(百分比欄位)]
    from itertools import chain
    最大值顯著欄位.extend(整數欄位)
    最大值顯著欄位.extend(實數欄位)
    最大值顯著欄位.extend(百分比欄位)
    最大值顯著欄位 = [*set(最大值顯著欄位)]
    最大值顯著欄位 = [x for x in 最大值顯著欄位 
                         if x not in 百分比漸層欄位 
                         or x not in 不排序欄位]
    s = df.style.pipe(標準格式(整數欄位, 實數欄位, 百分比欄位
                              ,日期欄位
                              ,最大值顯著欄位
                              ,隱藏欄位
                              ,百分比漸層欄位
                              ,採用民國日期格式=採用民國日期格式
                              ))
    tp = df.copy()
    for c in df.columns: tp.loc[c] = c

    if 除錯提示:
        from pathlib import Path
        html = Path.home() / 'TEMP' / 'output.html'
        tp.to_html(html)
        import os
        os.system(f'start {html}')

    if 顯示提示:
        s = s.set_tooltips(tp)

    if 顯示:
        from pathlib import Path
        html = Path.home() / 'TEMP' / 'showdf.html'
        html.parent.mkdir(parents=True, exist_ok=True)
        s.to_html(html, encoding='utf8', doctype_html=True)
        import os
        os.system(f'start {html}')
    return s

def read_csv(*args, **kwargs):
    '編碼錯誤預設使用 replace。'
    from pathlib import Path
    fn = args[0]
    if isinstance(fn, Path):
        fn = str(fn)
    with open(fn, encoding=kwargs['encoding'], errors='replace') as f:
        import pandas as pd
        df = pd.read_csv(f, *args[1:], **kwargs)
        return df

def read_fwf(*args, **kwargs): 
    import pandas as pd
    mbyte_columns = kwargs['mbyte_columns']
    names = kwargs['names']
    encoding = kwargs['encoding']
    del kwargs['mbyte_columns']
    del kwargs['encoding']
    df = pd.read_fwf(args[0], header=None
                    ,converters={h:str for h in names}
                    ,encoding = 'cp437' # 利用cp437剛好用1位元編碼特性，表達ByteString
                    ,**kwargs
                    )
    def convert(cp437:str):
        try:
            return cp437.encode('cp437').decode(encoding, errors='replace') 
        except AttributeError:
            return cp437
    for c in mbyte_columns: 
        df[c] = df[c].map(convert)
    return df

def read_docx_tables(filename, tab_id=None, **kwargs):
    from warnings import warn
    warn("警告：【read_docx_tables】函數將廢棄，請改用【read_docx】函數！")
    return read_docx_tables(filename, tab_id, **kwargs)

def read_docx(filename, tab_id=None, **kwargs):
    """
    讀取 Word 文件(.docx)表格至 Pnadas DataFrame
    parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)

    Parameters:
        filename:   file name of a Word Document

        tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                    When [None] - return a list of DataFrames (parse all tables)

        kwargs:     arguments to pass to `pd.read_csv()` function

    Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
    """
    from docx import Document
    def read_docx_tab(tab, **kwargs):
        import io
        import csv
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        import pandas as pd
        return pd.read_csv(vf, **kwargs)

    doc = Document(filename)
    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
    else:
        try:
            return read_docx_tab(doc.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise

def 分割鏈(鏈, 長度):
    from itertools import islice
    it = iter(鏈)
    size = 長度
    chunk = list(iter(lambda: list(islice(it, size)), []))
    return chunk

def 製作排行榜(排行個數或名次清單, 排行鍵, 漸增=False):
    def _製作排行榜(資料函數):
        from functools import wraps
        @wraps(資料函數)
        def 排行榜(*args, **kargs):
            df = 資料函數(*args, **kargs)
            df.sort_values(排行鍵, ascending=漸增, inplace=True)
            df.reset_index(drop=True, inplace=True)
            df.index = df.index+1
            df.index.name = '名次'
            if isinstance(排行個數或名次清單, int):
                排行個數 = 排行個數或名次清單 
                df = df.iloc[:排行個數]
            elif isinstance(排行個數或名次清單, list):
                名次清單 = 排行個數或名次清單 
                df = df.iloc[map(lambda n: n-1, 名次清單)]
            return df
        return 排行榜
    return _製作排行榜

def 連續增減次數(時序資料):
    資料名 = 時序資料.name
    df = 時序資料.to_frame()
    df['index']= df.index
    當期值 = df[資料名] 
    前期值 = 當期值.shift()
    環比 = 當期值 - 前期值
    消長 = 環比/環比.abs()
    變化 = 消長 != 消長.shift()
    df['分組']= 變化.cumsum()
    連續次數 = df.groupby('分組')['index'].rank(method='first')
    連續次數 = 連續次數 * 消長
    return 連續次數 

def 清查資料表非全數值欄(df, 列出非數值明細之欄=None):
    import pandas as pd

    # 將所有值轉換為數字類型，無法轉換的值設置為 NaN
    df_converted = df.apply(pd.to_numeric, errors='coerce')

    # 找出無法轉換的值
    invalid_cells = df_converted.isna()

    # 遍歷所有無法轉換的值並輸出
    for column in df.columns:
        if invalid_cells[column].any():
            invalid_values = df[column][invalid_cells[column]]
            logger.warn(f"{column} 欄有無法轉換成數值之資料值！")
            if column in 列出非數值明細之欄:
                logger.warn(f"非數值之資料值及其索引: {invalid_values.index.tolist()} - {invalid_values.tolist()}")
            else:
                logger.debug(f"非數值之資料值及其索引: {invalid_values.index.tolist()} - {invalid_values.tolist()}")


def 重名加序(columns):
    from zhongwen.表 import 重名加序
    return 重名加序(columns)
