from lark import Lark, Transformer
from pathlib import Path
import logging
logger = logging.getLogger(Path(__file__).stem)

TITLE = -1
PROPOSER = -2
EVENTS = -3

def get_doc_text(doc):
    "取得舊版 Word 檔案 .doc 內文"
    import win32com.client as win32
    import os
    import pyperclip

    # 啟動 Word 應用程式
    word = win32.gencache.EnsureDispatch('Word.Application')
    word.Visible = False  # 讓 Word 在背景執行
    try:
        # 打開文件
        doc = word.Documents.Open(str(doc))
        return f"源檔：{doc}\n{doc.Content.Text}"
    except Exception as e:
        print(f"發生錯誤：{e}")
    finally:
        # 關閉文件和 Word 程式，確保資源被釋放
        if 'doc' in locals() and doc:
            doc.Close(False)  # False 表示不儲存變更
        word.Quit()
        
def 列印(docs):
    from collections.abc import Iterable 
    import win32com.client as win32
    from collections.abc import Iterable 
    if isinstance(docs, str) or not isinstance(docs, Iterable):
        docs = [docs]
    word = win32.Dispatch("Word.Application")
    for docfile in docs:
        print(f'列印{docfile}')
        doc = word.Documents.Open(docfile)
        doc.PrintOut()
        doc.Close(False)
    word.Quit()

def read_docx(doc):
    '讀取word文件文字'
    from docx import Document
    doc = Document(doc)
    return "\n".join([p.text for p in doc.paragraphs])

def fread_docx(doc):
    from file import 最新符合檔 
    doc = 最新符合檔(Path("D:\\g"), f"**/*{doc}*.docx")
    return read_docx(doc)

def 轉換游標標號層次(層級):
    '游標處符合模式字串，否傳回None'
    import vim
    cb = vim.current.buffer
    l = vim.current.window.cursor[0] - 1
    line = cb[l]
    line = 設定標號層次(層級, line)
    cb[l] = line

def 游標模式替換(模式, 替換函式):
    '游標處符合模式字串，否傳回None'
    import vim
    import re
    cb = vim.current.buffer
    _a, l, c, _a, _a = map(lambda s: int(s)-1, vim.eval('getcursorcharpos()'))
    line = cb[l] 
    p = re.compile(模式)
    i = 0
    while m:=p.search(line, i):
        if m.start() <= c < m.end():
            替換 = 替換函式(m[0]) 
            cb[l] = line[:m.start()] + 替換 + line[m.end():]
            return [m.group(0), m.start(), m.end(), line, cb, l]
        i = m.end()+1
    raise ValueError(f'游標處無此模式，模式:{模式}，游標:{c}，行:{line}')

def 轉文字檔(filepath):
    from pathlib import Path
    import os
    output_txt = Path(filepath).with_suffix('.txt')
    cmd = f'pandoc -t plain "{filepath}" -o "{output_txt}"'
    os.system(cmd)

def 複製文字(filepath):
    '複製文字至剪貼簿'
    import clipboard
    with open(args.copy_text, 'r', encoding='utf8') as f:
        clipboard.copy(f.read())
    print(clipboard.paste())

def 設定環境():
    from zhongwen.winman import 增加檔案右鍵選單功能, 建立傳送到項目, where
    from zhongwen.office_document import 設定微軟辦公室軟體共用範本
    from shutil import copy
    import sys
    設定微軟辦公室軟體共用範本()

    cmd =  f'"{sys.executable}" -m zhongwen.office_document --print %* && pause'
    建立傳送到項目('列印微軟文件', cmd)

    cmd =  f'"{sys.executable}" -m zhongwen.office_document --to_xlsx %* && pause'
    建立傳送到項目('2xlsx', cmd)

    cmd = f'{sys.executable} -m zhongwen.office_document --doc2pdf "%1"' 
    增加檔案右鍵選單功能('2pdf', cmd, 'Word.Document.8')  # .doc
    增加檔案右鍵選單功能('2pdf', cmd, 'Word.Document.12') # .docx

    cmd = f'{sys.executable} -m zhongwen.office_document --doc2docx "%1"' 
    增加檔案右鍵選單功能('2docx', cmd, 'Word.Document.8')

    cmd = f'{sys.executable} -m zhongwen.office_document --to_text "%1"' 
    增加檔案右鍵選單功能('2txt', cmd, 'Word.Document.8')
    增加檔案右鍵選單功能('2txt', cmd, 'Word.Document.12') # .docx

    cmd = f'{sys.executable} -m zhongwen.office_document --copy_text "%1"' 
    增加檔案右鍵選單功能('複製文字', cmd, 'Word.Document.8') # .docx
    增加檔案右鍵選單功能('複製文字', cmd, 'Word.Document.12') # .docx

    cmd = f'{sys.executable} -m zhongwen.office_document --save_highlight "%1"' 
    增加檔案右鍵選單功能('另存醒目文字', cmd, 'Word.Document.8') # .docx
    增加檔案右鍵選單功能('另存醒目文字', cmd, 'Word.Document.12') # .docx
    
    cmd = f'{sys.executable} -m zhongwen.office_document --md2docx "%1" && pause' 

    增加檔案右鍵選單功能('markdown2docx', cmd, '.md') # .docx
    gvim = where('gvim')[-1]
    cmd = rf'"{gvim}" "%1"' 
    增加檔案右鍵選單功能('open', cmd, '.md') # .docx

def setup_normal_dotm_dir():
    '設定 resource 為 Word 預設範本資料夾'
    import winreg
    import os

    # Normal.dotm 的資料夾
    target_path = Path(__file__).parent / 'resource'

    # 支援的 Office 版本（16.0=Office 2016/2019/2021/365，15.0=2013，14.0=2010）
    office_versions = ["16.0", "15.0", "14.0"]

    # 依序檢查目前版本
    found_version = None
    for version in office_versions:
        reg_path = fr"Software\Microsoft\Office\{version}\Word\Options"
        try:
            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, reg_path, 0, winreg.KEY_ALL_ACCESS) as key:
                found_version = version
                break
        except FileNotFoundError:
            continue

    if found_version:
        reg_path = fr"Software\Microsoft\Office\{found_version}\Word\Options"
        with winreg.OpenKey(winreg.HKEY_CURRENT_USER, reg_path, 0, winreg.KEY_ALL_ACCESS) as key:
            winreg.SetValueEx(key, "USER TEMPLATES", 0, winreg.REG_SZ, str(target_path))
        print(f"已將 Word Normal.dotm 目錄設定為：{target_path}")
        print("請重新啟動 Word 後生效。")
    else:
        print("未找到已安裝的 Office Word 版本，請確認 Office 是否已安裝。")

def 合併文件(文件集, 合併文件檔=None, 文件以換頁符號分隔=True):
    '合併微軟辦公室文件(Word)，僅副檔名為 docx 者。' 
    from docxcompose.composer import Composer
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from pathlib import Path
    文件集 = sorted(文件集)

    if not 合併文件檔:
        合併文件檔 = Path(文件集[0]).parent / '合併文件.docx'
    
    input_paths = 文件集
    # Create a new document based on the first input document
    merged_document = Document(input_paths[0])
    composer = Composer(merged_document)

    def 插入換頁符號(document):
        # Create a page break element
        page_break = OxmlElement('w:br')
        page_break.set(qn('w:type'), 'page')
        # Add the page break to the last paragraph
        document.paragraphs[-1]._p.append(page_break)

    for input_path in input_paths[1:]:
        if 文件以換頁符號分隔:
            插入換頁符號(merged_document)
        # Append each subsequent document
        composer.append(Document(input_path))

    # Save the merged document to the specified output path
    composer.save(str(合併文件檔))

def doc2docx(docs, outputdir=None):
    '微軟辦公室文件為 doc 格式者，轉成 docx 格式。'
    from collections.abc import Iterable 
    import comtypes.client
    import pypandoc
    import os
    if isinstance(docs, str) or not isinstance(docs, Iterable):
        docs = [docs]
    for doc in docs:
        input_path = str(doc)
        if outputdir:
            output_path = str(Path(outputdir) / Path(doc).with_suffix('.docx').name)
        else:
            output_path = str(Path(doc).with_suffix('.docx'))
        # Ensure the input file is a .doc file
        if not input_path.lower().endswith('.doc'):
            raise ValueError("Input file must be a .doc file")
        # Ensure the output file is a .docx file
        if not output_path.lower().endswith('.docx'):
            raise ValueError("Output file must be a .docx file")

        # Create a Word application object
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False

        # Open the .doc file
        doc = word.Documents.Open(input_path)

        # Save the file as .docx
        doc.SaveAs(output_path, FileFormat=16)  # 16 corresponds to the wdFormatDocumentDefault format

        # Close the document and quit the application
        doc.Close()
        word.Quit()

def doc2pdf(words, output_dir=None):
    from collections.abc import Iterable 
    from pathlib import Path
    import win32com.client
    if isinstance(words, str) or not isinstance(words, Iterable):
        words = [words]
    word = win32com.client.Dispatch('Word.Application')
    for w in words:
        w = Path(w)
        if output_dir:
            pdf = output_dir / w.with_suffix('.pdf').name
        else:
            pdf = w.with_suffix('.pdf')
        doc = word.Documents.Open(str(w))
        doc.SaveAs(str(pdf), FileFormat=17)
        doc.Close()
    word.Quit()

def 標題階層編號轉中文編號(標題):
    from zhongwen.number import 中文數字, 大寫中文數字
    import re
    s = 標題
    pat = r'((\d+.)*(\d+)[\t\s])(.*)'
    if m:=re.match(pat, s):
        階層 = m[1].count('.')
        編號 = m[3]
        內容 = m[4]
        if 階層==0:
            編號 = ''
        elif 階層==1:
            編號 = f'{大寫中文數字(編號)}、'
        elif 階層==2: 
            編號 = f'{中文數字(編號)}、'
        elif 階層==3: 
            編號 = f'({中文數字(編號)})'
        elif 階層==4: 
            編號 = f'{編號}.'
        elif 階層==5: 
            編號 = f'({編號})'
        return 編號 + 內容
    return s

def html2docx(html):
    from pathlib import Path
    import win32com.client
    import pypandoc
    import os
    output_path = "output.docx"
    html = Path(html)
    docx = html.with_suffix('.docx')
    pypandoc.convert_text(html.read_text(encoding='utf8'), "docx", format="html", outputfile=docx)
    logger.info(f'{html.name}已轉成{docx.name}！')
    return docx


def markdown2docx(md):
    from zhongwen.數 import 取中文數字, 取大寫中文數字
    from docx import Document
    from pathlib import Path
    import re
    import os
    md = Path(md)
    docx = Path(__file__).parent / md.with_suffix('.docx')
    temp = r'c:\GitHub\zhongwen\zhongwen\resource\審核報告範本.docx'
    cmd = f'pandoc -f markdown+east_asian_line_breaks -t docx '
    cmd += f'--reference-doc="{temp}" --number-sections '
    cmd += f'-o "{docx}" {md}'
    os.system(cmd)
    document = Document(str(docx))
    階層 = 0

    def 取中文階層編號(編號, 階層):
        if 階層==0:
            return ''
        elif 階層==1:
            return f'{取大寫中文數字(編號)}、'
        elif 階層==2: 
            return f'{取中文數字(編號)}、'
        elif 階層==3: 
            return f'({取中文數字(編號)})'
        elif 階層==4: 
            return f'{編號}.'
        elif 階層==5: 
            return f'({編號})'
        return 編號

    for paragraph in document.paragraphs:
        if m:='Heading' in paragraph.style.name:
            runs = list(paragraph.runs)
            if len(runs) > 0 and number_run=runs[0]:
                text = number_run.text
                pat = r'^((\d+.)*(\d+))$'
                if m:=re.match(pat, text):
                    階層 = m[1].count('.')
                    編號 = 取中文階層編號(m[3])
                    number_run.text = text.replace(m[1], 編號)
        if 'Normal' in paragraph.style.name:
            if int(階層)>0:
                paragraph.style = f'內文{階層+1}'
    f = Path(f)
    nf = f.with_stem(f'{f.stem}新')
    document.save(str(nf))
    os.system(f'start {nf}')

def 另存醒目文字(docx):
    from docx import Document
    from pathlib import Path
    docx = Path(docx)
    doc = Document(str(docx))

    highlighted_texts = []

    for para in doc.paragraphs:
        本段落前累積高亮區塊數 = len(highlighted_texts)
        for run in para.runs:
            if run.font.highlight_color:  
                highlighted_texts.append(run.text)
        if len(highlighted_texts) > 本段落前累積高亮區塊數:
            highlighted_texts.append('\n')

    highlighted_texts = ''.join(highlighted_texts )
    highlighted_texts = f'宣讀{highlighted_texts }'
    print(highlighted_texts)
    txt = docx.with_suffix('.txt')
    txt.write_text(highlighted_texts, encoding='utf8')

def to_xlsx(odses_or_pdfs):
    from collections.abc import Iterable 
    from pathlib import Path
    from zhongwen.pdf import to_excel as pdf2xlsx
    import pandas as pd
    import pyexcel as p
    odses = odses_or_pdfs
    if isinstance(odses, str) or not isinstance(odses, Iterable):
        odses = [odses]
    odses = [Path(ods) for ods in odses]
    for ods in odses:
        xlsx = ods.with_suffix('.xlsx')
        try:
            if ods.suffix == '.ods':
                p.save_book_as(file_name=str(ods), dest_file_name=str(xlsx))
            elif ods.suffix == '.pdf':
                pdf2xlsx(ods)
            logger.info(f"{ods.name}->{xlsx.name}")
        except Exception as e:
            logger.info(f"{ods.name}轉換失敗：{e}")

if __name__ == '__main__':
    import argparse
    import clipboard
    logging.basicConfig(level=logging.INFO)
    parser = argparse.ArgumentParser()
    parser.add_argument("--setup", help="設定環境", action="store_true")
    parser.add_argument('--print', nargs='+', type=str, help='列印微軟文件')
    parser.add_argument("--setup_template", help="設定 resource 為 Word 預設範本資料夾"
                       ,action="store_true")
    parser.add_argument("--doc2pdf", type=str, help="轉 PDF 格式。", required=False)
    parser.add_argument("--doc2docx", type=str, help="doc 轉 docx 格式。", required=False)
    parser.add_argument("--to_xlsx", type=str, help="ods 或 pdf 轉 xlsx 格式。", required=False)
    parser.add_argument("--md2docx", type=str, help="markdown 轉 docx 格式。", required=False)
    parser.add_argument("--to_text", type=str, help="轉文字檔", required=False)
    parser.add_argument("--copy_text", type=str, help="複製文字", required=False)
    parser.add_argument("--level_number_to_chinese_number", help="標題階層編號轉中文編號")
    parser.add_argument("--save_highlight", help="另存醒目文字")

    args = parser.parse_args()
    if args.setup:
        設定環境()
    elif docs := args.print:
        列印(docs)
    elif args.setup_template:
        setup_normal_dotm_dir()
    elif args.doc2pdf:
        doc2pdf(args.doc2pdf)
    elif args.doc2docx:
        doc2docx(args.doc2docx)
    elif args.to_xlsx:
        to_xlsx(args.to_xlsx)
    elif args.md2docx:
        markdown2docx(args.md2docx)
    elif args.to_text:
        轉文字檔(args.to_text)
    elif args.copy_text:
        複製文字(args.copy_text)
    elif args.level_number_to_chinese_number:
        print(標題階層編號轉中文編號(args.level_number_to_chinese_number))
    elif args.save_highlight:
        print(另存醒目文字(args.save_highlight))
